# -*- coding: utf-8 -*-
"""生成软件说明书 Word 文档"""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

ASSETS = r"d:\Project\DeltaRent\docs\report\assets"
OUTPUT = r"d:\Project\DeltaRent\软件说明书.docx"

doc = Document()

# ── 全局样式 ──
style = doc.styles["Normal"]
style.font.name = "宋体"
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

for level, (sz, name) in enumerate([(16, "黑体"), (14, "黑体"), (12, "黑体")], 1):
    h = doc.styles[f"Heading {level}"]
    h.font.name = name
    h.font.size = Pt(sz)
    h.font.bold = True
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.element.rPr.rFonts.set(qn("w:eastAsia"), name)

# ── 页边距 ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ── 辅助函数 ──
def add_para(text, bold=False, align=None, size=None, font_name=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = size
    if font_name:
        run.font.name = font_name
        run.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    return p

def add_number(text):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.space_after = Pt(4)
    return p

def add_image(filename, width_inches=5.0, caption=None):
    img_path = os.path.join(ASSETS, filename)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches))
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(12)
            for r in cap.runs:
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(102, 102, 102)

def make_table(headers, rows, col_widths=None):
    ncols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        # shading
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D5E8F0" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    # data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(11)
    if col_widths:
        for ri, row in enumerate(tbl.rows):
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()  # spacing after table
    return tbl

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

# ================================================================
# 封面
# ================================================================
for _ in range(4):
    doc.add_paragraph()

add_para("湘潭大学", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(26), font_name="黑体")
add_para("2026 信息系统实践课程", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(22), font_name="黑体")
doc.add_paragraph()
add_para("软件说明书", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(22), font_name="黑体")
add_para("（结项版）", align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(16), font_name="黑体")

for _ in range(4):
    doc.add_paragraph()

for line in [
    "项目名称：三角洲行动账号租赁管理系统（DeltaRent）",
    "学　　院：计算机学院",
    "小组序号：18",
    "成员姓名：罗季维 / 朱江南 / 李彬菲 / 胡万军",
    "指导老师：尹兆远",
    "更新日期：2026年5月19日",
]:
    add_para(line, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14))

doc.add_page_break()

# ================================================================
# 基本信息
# ================================================================
doc.add_heading("基本信息", level=1)
for label, val in [
    ("项目名称", "三角洲行动账号租赁管理系统（DeltaRent）"),
    ("学　　院", "计算机学院"),
    ("小组序号", "18"),
    ("成员姓名", "罗季维 / 朱江南 / 李彬菲 / 胡万军"),
    ("指导老师", "尹兆远"),
    ("当前版本", "结项版"),
    ("更新日期", "2026年5月19日"),
]:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}：")
    r1.bold = True
    p.add_run(val)

divider()

# ================================================================
# 一、项目概述
# ================================================================
doc.add_heading("一、项目概述", level=1)

doc.add_heading("1. 项目背景", level=2)
add_para(
    "“三角洲行动”相关玩家围绕账号租赁、账号展示、订单咨询、售后申诉等场景，"
    "存在较高频的信息撮合需求。现有交易方式多依赖微信群、QQ群和人工客服完成，存在以下问题："
)
add_bullet("账号资源描述不统一，用户难以快速比较账号价值")
add_bullet("下单、确认、交付、售后等流程依赖人工同步，效率较低")
add_bullet("订单、用户、申诉和公告信息缺乏统一沉淀")
add_bullet("后台运营缺乏统一管理界面和统计视图")
add_para(
    "本项目围绕“账号展示 + 租赁下单 + 订单中心 + 后台运营”设计并实现了一个前后端分离的"
    "网站系统原型，用于完成课程要求中的需求分析、系统设计、编码实现、测试验证与文档交付。"
)

doc.add_heading("2. 系统目标", level=2)
add_para("本系统的最终目标如下：")
add_bullet("实现一个可运行、可演示、业务链路完整的账号租赁管理系统原型")
add_bullet("支持注册登录、账号浏览、账号上架、订单创建、订单查看、个人中心、公告展示等前台核心功能")
add_bullet("支持后台看板、用户管理、账号管理、订单管理、公告管理等运营功能")
add_bullet("采用前后端分离架构，形成清晰的页面层、接口层、业务层和数据层结构")
add_bullet("在满足课程要求的前提下，为后续接入短信、支付、消息队列、对象存储等能力预留扩展空间")

doc.add_heading("3. 开发环境", level=2)
make_table(
    ["类别", "当前方案"],
    [
        ["前端", "Vue 3 + TypeScript + Vite + Vue Router + Pinia + Element Plus"],
        ["后端", "Spring Boot 3 + Spring Security + JWT + MyBatis"],
        ["数据库", "MySQL 8.0"],
        ["缓存/扩展预留", "Redis"],
        ["构建工具", "npm、Gradle Wrapper"],
        ["运行环境", "Node.js 24、JDK 21、Windows 11"],
        ["版本管理", "Git + GitHub"],
        ["远端仓库", "https://github.com/JWLuo0719/DeltaRent"],
    ],
    col_widths=[4, 12],
)

divider()

# ================================================================
# 二、需求分析
# ================================================================
doc.add_heading("二、需求分析", level=1)

doc.add_heading("1. 功能需求", level=2)

doc.add_heading("1.1 目标用户", level=3)
make_table(
    ["角色", "主要诉求", "典型使用方式"],
    [
        ["游客", "查看账号资源、浏览公告、了解平台规则", "浏览首页、租号大厅、公告信息"],
        ["普通用户", "注册登录、账号上架、租赁下单、查看订单、提交申诉", "登录后浏览账号详情、创建订单、发布账号、查看订单进度"],
        ["客服", "处理订单、维护账号状态、协助售后", "进入后台处理订单和账号管理"],
        ["管理员", "管理用户、公告、账号、订单和统计数据", "进入后台管理台查看看板并执行运营管理操作"],
    ],
    col_widths=[3, 6, 7],
)

doc.add_heading("1.2 前台功能需求", level=3)
make_table(
    ["模块", "主要功能"],
    [
        ["首页门户", "平台介绍、公告展示、热点账号、业务入口"],
        ["用户认证", "注册、登录、找回密码、验证码发送"],
        ["租号大厅", "账号列表展示、关键词搜索、标签筛选、排序、分页浏览"],
        ["账号详情/下单", "查看账号价值信息、选择租赁时长、填写联系方式并提交订单"],
        ["我要上架", "提交账号信息、自动计算租金、进入待审核状态"],
        ["订单中心", "查看我的订单、按状态筛选、查看订单详情、取消订单、申请售后"],
        ["个人中心", "查看资料、修改昵称、修改密码、退出登录"],
    ],
    col_widths=[4, 12],
)

doc.add_heading("1.3 后台功能需求", level=3)
make_table(
    ["模块", "主要功能"],
    [
        ["数据看板", "查看待确认订单、进行中订单、注册用户数、最近订单"],
        ["用户管理", "查看用户列表、按手机号/角色/状态筛选、修改角色与状态"],
        ["账号管理", "新增账号、编辑账号、上下架、删除账号"],
        ["订单管理", "查看全部订单、按状态过滤、修改订单状态"],
        ["公告管理", "查看公告、新增公告、修改公告、删除公告"],
        ["售后处理", "查看用户申诉、客服处理申诉结果"],
    ],
    col_widths=[4, 12],
)

doc.add_heading("1.4 主要业务流程", level=3)

add_para("业务流程一：账号租赁", bold=True)
add_number("用户注册或登录系统")
add_number("浏览租号大厅，按关键词、标签、状态或价格筛选账号")
add_number("查看账号详情与资源说明，选择租赁时长并提交订单")
add_number("订单进入“待确认”状态")
add_number("客服或管理员在后台确认订单并推进状态流转")
add_number("用户在订单中心查看订单详情、时间线和状态变化")
add_number("租赁结束后进入已完成状态，若存在问题可发起售后申诉")

add_para("业务流程二：账号上架", bold=True)
add_number("登录用户进入“我要上架”页面")
add_number("填写账号标题、上号方式、哈夫币数额、比例、押金、段位、红皮等字段")
add_number("系统根据比例自动计算租金")
add_number("用户提交后账号进入待审核/维护状态")
add_number("管理员或客服在后台审核后再决定上架到租号大厅")

add_para("业务流程三：后台运营", bold=True)
add_number("管理员登录后台进入数据看板")
add_number("管理账号信息、公告内容、用户状态和订单状态")
add_number("客服在权限范围内处理账号与订单相关事务")
add_number("系统通过统计视图辅助日常运营管理")

doc.add_heading("2. 非功能需求", level=2)

doc.add_heading("2.1 性能要求", level=3)
add_bullet("页面和接口在课程演示环境下应保持秒级响应")
add_bullet("列表查询支持分页与筛选，避免一次性加载过大数据集")
add_bullet("前端资源能够稳定打包构建，后端能够稳定启动并提供 REST 接口")

doc.add_heading("2.2 安全要求", level=3)
add_bullet("使用 Spring Security + JWT 实现基于令牌的身份认证")
add_bullet("使用 RBAC 角色控制区分游客、普通用户、客服、管理员的访问范围")
add_bullet("密码采用 BCrypt 加密存储，不以明文形式保存")
add_bullet("后端对关键接口做权限限制，例如后台用户管理、订单管理、公告管理等")

doc.add_heading("2.3 兼容性要求", level=3)
add_bullet("支持 Chrome、Edge 等主流现代浏览器")
add_bullet("前台兼顾桌面端与常规移动端浏览")
add_bullet("后台以桌面端操作体验为主")

doc.add_heading("2.4 可维护性要求", level=3)
add_bullet("前后端分离，接口集中管理，便于并行开发与联调")
add_bullet("业务模块按“认证、账号、订单、公告、用户、后台管理”拆分")
add_bullet("数据库脚本单独维护，便于初始化和后续迭代")

divider()

# ================================================================
# 三、系统设计
# ================================================================
doc.add_heading("三、系统设计", level=1)

doc.add_heading("1. 系统架构", level=2)
add_para(
    "系统采用“Vue 前端 + Spring Boot 后端 + MySQL 数据库”的前后端分离架构。"
    "前台和后台共用一套后端接口体系，通过角色权限限制不同功能入口。"
)
add_image("fig_architecture.png", 5.0, "图 1  系统架构图")

add_para("架构说明：")
add_bullet("前端负责页面展示、表单校验、页面路由控制和接口调用")
add_bullet("后端负责认证授权、业务规则处理、数据持久化和接口输出")
add_bullet("MySQL 存储用户、账号、订单、公告、申诉、日志等核心数据")
add_bullet("Redis 作为缓存与验证码能力预留，当前版本优先完成数据库主链路")

doc.add_heading("2. 模块设计", level=2)
make_table(
    ["模块", "说明", "对应实现"],
    [
        ["认证模块", "登录、注册、找回密码、验证码发送、JWT 生成与校验", "AuthController、LoginView.vue、RegisterView.vue"],
        ["账号模块", "账号列表、详情、筛选、后台账号管理、用户上架账号", "RentalController、RentalListView.vue、ProductManageView.vue"],
        ["订单模块", "创建订单、我的订单、订单详情、取消订单、后台订单管理", "OrderController、OrderListView.vue、OrderManageView.vue"],
        ["用户模块", "个人资料查看与修改、密码修改、后台用户管理", "UserController、AdminUserController、ProfileView.vue"],
        ["公告模块", "前台公告展示、后台公告 CRUD", "NoticeController、NoticeManageView.vue"],
        ["售后模块", "申诉提交与后台处理", "AppealController"],
        ["看板模块", "指标概览、最近订单、后台主入口", "DashboardController、StatsView.vue、AdminDashboardView.vue"],
    ],
    col_widths=[3, 6, 7],
)

doc.add_heading("3. 数据库设计", level=2)

doc.add_heading("3.1 E-R 关系说明", level=3)
add_para("系统数据库包含以下主要实体及其关系：")
add_bullet("用户（SYS_USER）与角色（SYS_ROLE）通过用户角色关联表（SYS_USER_ROLE）实现多对多关系")
add_bullet("用户可以创建多个租赁订单（RENTAL_ORDER）")
add_bullet("账号商品（RENTAL_PRODUCT）可以被多个订单引用")
add_bullet("用户可以提交多个售后申诉（APPEAL_RECORD）")
add_bullet("用户可以发布多个公告（NOTICE）")
add_bullet("用户操作记录在操作日志表（OPERATION_LOG）中")

doc.add_heading("3.2 主要数据表", level=3)
make_table(
    ["表名", "说明"],
    [
        ["sys_user", "用户基础信息，含用户名、手机号、密码哈希、状态、密码更新时间等"],
        ["sys_role", "角色定义，如 ADMIN、USER、CS"],
        ["sys_user_role", "用户与角色的关联表"],
        ["rental_product", "账号商品表，保存价格、标签、哈夫币数量、段位、皮肤等字段"],
        ["rental_order", "租赁订单表，保存订单号、用户、商品、时长、金额、状态、交付备注等"],
        ["notice", "公告内容表"],
        ["sms_verify_code", "短信验证码表，用于找回密码流程"],
        ["appeal_record", "售后申诉记录表"],
        ["operation_log", "操作日志表"],
    ],
    col_widths=[4, 12],
)

doc.add_heading("3.3 核心表设计摘要", level=3)

add_para("用户表 sys_user", bold=True)
make_table(
    ["字段", "类型", "说明"],
    [
        ["id", "BIGINT", "主键"],
        ["username", "VARCHAR(50)", "用户名，唯一"],
        ["phone", "VARCHAR(20)", "手机号，唯一"],
        ["password_hash", "VARCHAR(255)", "BCrypt 加密密码"],
        ["nickname", "VARCHAR(50)", "昵称"],
        ["status", "TINYINT", "用户状态"],
        ["password_updated_at", "DATETIME", "密码更新时间"],
        ["deleted_at", "DATETIME", "软删除字段"],
    ],
    col_widths=[5, 4, 7],
)

add_para("账号表 rental_product", bold=True)
make_table(
    ["字段", "类型", "说明"],
    [
        ["id", "BIGINT", "主键"],
        ["name", "VARCHAR(100)", "账号标题"],
        ["category", "VARCHAR(50)", "分类"],
        ["tag_text", "VARCHAR(255)", "标签"],
        ["hour_price", "DECIMAL(10,2)", "小时单价"],
        ["coin_amount", "BIGINT", "哈夫币数量"],
        ["rank_text / character_skin_text / knife_skin_text", "VARCHAR", "三角洲账号扩展字段"],
        ["status", "VARCHAR(30)", "可租状态"],
    ],
    col_widths=[5, 4, 7],
)

add_para("订单表 rental_order", bold=True)
make_table(
    ["字段", "类型", "说明"],
    [
        ["id", "BIGINT", "主键"],
        ["order_no", "VARCHAR(40)", "订单号，唯一"],
        ["user_id", "BIGINT", "下单用户"],
        ["product_id", "BIGINT", "账号商品"],
        ["unit_price", "DECIMAL(10,2)", "单价"],
        ["rent_hours", "INT", "租赁时长"],
        ["order_amount", "DECIMAL(10,2)", "订单总金额"],
        ["contact_info", "VARCHAR(100)", "联系方式"],
        ["delivery_note", "VARCHAR(255)", "交付备注"],
        ["status", "VARCHAR(30)", "订单状态"],
        ["start_time / end_time", "DATETIME", "租赁开始与结束时间"],
    ],
    col_widths=[5, 4, 7],
)

doc.add_heading("3.4 初始数据", level=3)
add_para("当前 SQL 初始化脚本中已预置以下测试数据：")
add_bullet("3 个测试用户：管理员、普通用户、客服")
add_bullet("3 个测试角色：ADMIN、USER、CS")
add_bullet("多个测试账号商品")
add_bullet("测试公告")
add_para("数据库脚本位于 sql/schema_v2.sql。")

divider()

# ================================================================
# 四、系统实现
# ================================================================
doc.add_heading("四、系统实现", level=1)

doc.add_heading("1. 关键技术", level=2)

doc.add_heading("1.1 前端实现", level=3)
add_bullet("使用 Vue 3 + Composition API 实现页面逻辑拆分")
add_bullet("使用 Vue Router 管理游客页、用户页、后台页等多类路由")
add_bullet("使用动态导入优化路由加载")
add_bullet("使用 Pinia 管理登录态与用户信息")
add_bullet("使用 Axios 统一处理请求头注入和 401 响应跳转")
add_bullet("使用 Element Plus 搭建表单、表格、分页、弹窗等交互组件")

doc.add_heading("1.2 后端实现", level=3)
add_bullet("使用 Spring Boot 3 提供 RESTful API")
add_bullet("使用 Spring Security + JWT 完成登录鉴权与接口权限控制")
add_bullet("使用 MyBatis/Mapper 完成数据库访问")
add_bullet("使用 Gradle Wrapper 统一后端构建方式")
add_bullet("使用服务层封装订单、账号、用户、公告、申诉等业务规则")

doc.add_heading("1.3 关键实现点", level=3)

add_para("（1）JWT 登录与权限校验", bold=True)
add_para(
    "系统登录成功后返回包含用户 ID、用户名和角色的 JWT。前端通过请求拦截器自动附带 "
    "Authorization 请求头，后端基于 Spring Security 做访问控制。"
)
code = """http
    .csrf(AbstractHttpConfigurer::disable)
    .sessionManagement(session -> session.sessionCreationPolicy(SessionCreationPolicy.STATELESS))
    .authorizeHttpRequests(auth -> auth
        .requestMatchers("/api/health", "/api/auth/**", "/api/portal/**").permitAll()
        .requestMatchers(HttpMethod.GET, "/api/rentals/**", "/api/notices/**").permitAll()
        .requestMatchers("/api/dashboard/**").hasRole("ADMIN")
        .requestMatchers("/api/admin/**").hasRole("ADMIN")
        .requestMatchers("/api/appeals/**").authenticated()
        .anyRequest().authenticated()
    )
    .addFilterBefore(jwtAuthFilter, UsernamePasswordAuthenticationFilter.class);"""
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code)
run.font.name = "Consolas"
run.font.size = Pt(9)

add_para("（2）订单状态流转控制", bold=True)
add_para("订单模块支持创建订单、查看我的订单、查看订单详情、取消订单，以及后台修改订单状态。服务层对非法状态跳转做了限制。")
code2 = """RentalOrder order = orderService.create(userId, productId, rentHours, contactInfo, deliveryNote);
orderService.transitionStatus(order.getId(), "IN_PROGRESS");
orderService.transitionStatus(order.getId(), "COMPLETED");"""
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(code2)
run.font.name = "Consolas"
run.font.size = Pt(9)

add_para("（3）账号上架与租金自动计算", bold=True)
add_para(
    "“我要上架”页面支持输入账号三角洲专属字段，并根据“哈夫币数额 / 比例”"
    "自动计算租金，提交后默认进入待审核状态。"
)

add_para("（4）统一接口层封装", bold=True)
add_para("前端在 src/web/src/api/index.ts 中对登录、注册、账号列表、订单管理、后台管理等接口做了集中封装，便于联调与维护。")

doc.add_heading("2. 界面展示", level=2)
add_para("当前版本已完成并提交源码的主要页面如下：")

screenshots = [
    ("01-login.png", "图 2  登录页"),
    ("02-register.png", "图 3  注册页"),
    ("03-home-1.png", "图 4  首页门户页"),
    ("04-rentals.png", "图 5  租号大厅"),
    ("05-order-create.png", "图 6  创建订单页"),
    ("06-publish.png", "图 7  我要上架页"),
    ("07-orders.png", "图 8  我的订单页"),
    ("08-order-detail.png", "图 9  订单详情页"),
    ("09-profile.png", "图 10  个人资料页"),
    ("10-admin-dashboard.png", "图 11  后台数据看板"),
    ("11-admin-users.png", "图 12  后台用户管理页"),
    ("12-admin-accounts.png", "图 13  后台账号管理页"),
    ("13-admin-orders.png", "图 14  后台订单管理页"),
    ("14-admin-notice.png", "图 15  后台公告管理页"),
]
for fname, cap in screenshots:
    add_image(fname, 5.0, cap)

doc.add_heading("3. 核心代码片段", level=2)
add_para("当前版本的关键实现入口如下：")
add_bullet("登录注册接口：AuthController.java")
add_bullet("账号列表与账号管理接口：RentalController.java")
add_bullet("订单创建与订单管理接口：OrderController.java")
add_bullet("个人中心接口：UserController.java")
add_bullet("公告接口：NoticeController.java")
add_bullet("售后申诉接口：AppealController.java")
add_bullet("JWT 工具类：JwtUtil.java")

divider()

# ================================================================
# 五、系统测试
# ================================================================
doc.add_heading("五、系统测试", level=1)

doc.add_heading("1. 测试方案", level=2)

doc.add_heading("1.1 测试范围", level=3)
add_bullet("用户注册、登录、找回密码")
add_bullet("账号列表查询、筛选、浏览和详情查看")
add_bullet("账号上架流程")
add_bullet("订单创建、订单列表、订单详情、取消订单、售后申诉")
add_bullet("个人资料修改、密码修改")
add_bullet("后台用户、账号、订单、公告管理入口")
add_bullet("前后端项目构建与后端服务启动")

doc.add_heading("1.2 测试方法", level=3)
add_bullet("前端静态检查：vue-tsc")
add_bullet("前端构建测试：vite build")
add_bullet("后端服务层测试：gradlew test")
add_bullet("后端启动测试：访问 /api/health")
add_bullet("人工功能测试：基于真实后端逐模块执行测试用例")

doc.add_heading("2. 测试结果", level=2)

doc.add_heading("2.1 工程验证结果", level=3)
make_table(
    ["测试项", "执行方式", "结果"],
    [
        ["前端类型检查", "npm --prefix src/web run typecheck", "通过"],
        ["前端生产构建", "npm --prefix src/web run build", "通过"],
        ["后端服务层测试", "src/server/gradlew.bat test", "通过"],
        ["后端服务健康检查", "启动后访问 GET /api/health", "返回 status = UP"],
    ],
    col_widths=[4, 8, 4],
)

doc.add_heading("2.2 业务测试结果", level=3)
add_para("依据测试报告与测试用例文档，当前版本已完成 110 个功能测试用例，全部通过。")
make_table(
    ["模块", "用例数", "通过", "失败", "通过率"],
    [
        ["注册模块", "7", "7", "0", "100%"],
        ["登录模块", "8", "8", "0", "100%"],
        ["首页模块", "13", "13", "0", "100%"],
        ["租号大厅模块", "18", "18", "0", "100%"],
        ["我要上架模块", "8", "8", "0", "100%"],
        ["我的订单模块", "6", "6", "0", "100%"],
        ["个人中心模块", "9", "9", "0", "100%"],
        ["管理员模块", "29", "29", "0", "100%"],
        ["游客模块", "5", "5", "0", "100%"],
        ["找回密码模块", "7", "7", "0", "100%"],
        ["合计", "110", "110", "0", "100%"],
    ],
    col_widths=[4, 3, 3, 3, 3],
)

doc.add_heading("2.3 服务层测试覆盖点", level=3)
add_para("后端当前已实现服务层测试文件：")
add_bullet("OrderServiceTest：验证订单金额计算、不可租账号拦截、重复下单拦截、状态流转约束")
add_bullet("ProductServiceTest：验证账号筛选、排序、分类过滤、非法创建拦截")

doc.add_heading("3. 问题与改进", level=2)
make_table(
    ["当前问题", "说明", "后续改进"],
    [
        ["前端打包产物体积偏大", "vite build 存在 chunk size 警告", "后续通过更细粒度路由懒加载和公共包拆分优化"],
        ["界面截图资产仍不完整", "当前仓库主要保留架构图与测试文档", "如需正式归档，可补充登录页、后台页、订单页截图"],
        ["短信、支付、对象存储未接入真实外部服务", "当前为课程原型", "后续如继续开发，可替换为真实三方服务"],
    ],
    col_widths=[5, 5, 6],
)

divider()

# ================================================================
# 六、用户手册
# ================================================================
doc.add_heading("六、用户手册", level=1)

doc.add_heading("1. 安装部署说明", level=2)

doc.add_heading("1.1 环境要求", level=3)
add_bullet("Node.js 20 或更高版本")
add_bullet("JDK 21")
add_bullet("MySQL 8.0")
add_bullet("Windows PowerShell 或兼容终端")
add_para("后端默认连接本地 MySQL：")
for line in ["数据库：deltarent", "地址：localhost:3306", "账号：root", "密码：123456"]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(line)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
add_para("如需修改数据库配置，编辑 application.yml。")

doc.add_heading("1.2 数据库初始化", level=3)
add_para("在 MySQL 中执行：")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run("source sql/schema_v2.sql;")
run.font.name = "Consolas"
run.font.size = Pt(10)
add_para("也可以在 Windows 下执行：")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(".\\sql\\init.bat")
run.font.name = "Consolas"
run.font.size = Pt(10)

doc.add_heading("1.3 启动后端", level=3)
for line in ["cd src/server", ".\\gradlew.bat bootRun"]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(line)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
add_para("后端默认地址：http://localhost:8080")
add_para("健康检查：GET http://localhost:8080/api/health")

doc.add_heading("1.4 启动前端", level=3)
add_para("首次运行先安装依赖：")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run("npm --prefix src/web install")
run.font.name = "Consolas"
run.font.size = Pt(10)
add_para("启动开发服务器：")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run("npm --prefix src/web run dev")
run.font.name = "Consolas"
run.font.size = Pt(10)
add_para("前端默认地址：http://localhost:5173")

doc.add_heading("2. 操作指南", level=2)

doc.add_heading("2.1 普通用户使用流程", level=3)
add_number("访问首页或租号大厅")
add_number("注册并登录系统")
add_number("浏览账号列表，按标签、价格或状态筛选")
add_number("查看账号详情后创建订单")
add_number("在“我的订单”中查看订单状态和详情")
add_number("如需售后，可对订单发起申诉")
add_number("如有账号资源，也可进入“我要上架”页面提交账号信息")

doc.add_heading("2.2 管理员/客服使用流程", level=3)
add_number("使用管理员或客服账号登录")
add_number("进入后台管理页面")
add_number("查看看板指标和最近订单")
add_number("对账号状态、订单状态和公告内容执行管理操作")
add_number("管理员可进一步执行用户管理操作")

doc.add_heading("2.3 演示账号", level=3)
add_para("初始密码均为：123456")
make_table(
    ["角色", "手机号", "说明"],
    [
        ["管理员", "13800000000", "可访问后台用户、账号、订单、公告管理"],
        ["普通用户", "13900000000", "可浏览账号、创建订单、查看订单"],
        ["客服", "13700000000", "可处理账号与订单相关管理功能"],
    ],
    col_widths=[3, 4, 9],
)

divider()

# ================================================================
# 七、项目总结
# ================================================================
doc.add_heading("七、项目总结", level=1)

doc.add_heading("1. 成果总结", level=2)
add_para("截至结项阶段，项目已完成以下内容：")
add_bullet("完成项目计划书、中期与结项报告、测试文档等课程材料")
add_bullet("完成前后端分离项目结构搭建")
add_bullet("完成 MySQL 数据库建表、初始化与迁移脚本")
add_bullet("完成注册、登录、找回密码、JWT 登录态、权限控制等认证功能")
add_bullet("完成租号大厅、账号筛选、下单页、订单列表、订单详情、个人中心、我要上架等前台页面")
add_bullet("完成后台看板、用户管理、账号管理、订单管理、公告管理等后台页面")
add_bullet("完成对应的后端控制器、服务层与数据库访问逻辑")
add_bullet("完成 110 个功能测试用例和服务层测试，前后端可稳定构建运行")

doc.add_heading("2. 不足与改进方向", level=2)
add_para("当前系统已经满足课程结项展示要求，但仍有以下限制：")
add_bullet("当前系统仍是课程原型，不接入真实支付、真实短信网关和真实游戏自动化能力")
add_bullet("部分运营流程仍以手动审核和课程演示逻辑为主")
add_bullet("还可以继续补充更多自动化测试与前端性能优化")
add_bullet("若后续继续演进，可进一步接入对象存储、消息通知和完整审计日志")

doc.add_heading("3. 成员分工表", level=2)
make_table(
    ["姓名", "Git 账号", "主要职责", "对应产物/目录"],
    [
        ["罗季维", "JWLuo0719", "项目管理、需求分析、计划书、主文档整理、总体推进", "课程信息/、docs/plan/、docs/report/"],
        ["朱江南", "ZHU123-OK", "后端开发、数据库设计、认证与订单接口实现", "src/server/、sql/"],
        ["李彬菲", "LBF1105", "前端开发、页面实现、前后台联调、界面优化", "src/web/"],
        ["胡万军", "erousagi37", "测试用例整理、测试报告、文档协同整理", "docs/test/、docs/report/"],
    ],
    col_widths=[2.5, 3, 6, 4.5],
)

doc.add_heading("4. Git 提交记录", level=2)
add_para("远端仓库：https://github.com/JWLuo0719/DeltaRent")
add_para("结项阶段代表性提交如下：")
make_table(
    ["提交号", "提交说明"],
    [
        ["6c6f744", "Merge branch docs-test"],
        ["aef6d31", "完成全部测试用例和测试报告，通过率 100%"],
        ["3f681de", "完成个人中心模块测试用例"],
        ["c9b0f00", "更新我的订单模块测试用例"],
        ["7580e63", "补充我要上架模块审核流程测试用例"],
        ["ad5f1d0", "完成租号大厅和我要上架模块测试用例"],
        ["aa589e5", "优化订单列表筛选体验，修复租号/下单流程并完善布局"],
        ["64a72c3", "修复账号加载失败问题，重构数据库和前后端对应关系"],
    ],
    col_widths=[3, 13],
)

divider()

# ================================================================
# 附录
# ================================================================
doc.add_heading("附录", level=1)
add_bullet("课程模板：课程信息/2. 中期与结项报告.md")
add_bullet("系统架构图：docs/report/assets/fig_architecture.png")
add_bullet("测试报告：docs/test/测试报告.md")
add_bullet("测试用例：docs/test/测试用例.md")
add_bullet("数据库脚本：sql/schema_v2.sql")
add_bullet("源代码仓库：https://github.com/JWLuo0719/DeltaRent")

# ── 保存 ──
doc.save(OUTPUT)
fsize = os.path.getsize(OUTPUT)
print(f"软件说明书已生成: {OUTPUT}")
print(f"文件大小: {fsize / 1024:.2f} KB")
