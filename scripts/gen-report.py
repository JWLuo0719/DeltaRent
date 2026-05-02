"""Generate today's work report as a Word document."""
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page setup
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Title
title = doc.add_heading('DeltaRent Backend Development Daily Report', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run('Date: 2026-05-01').font.size = Pt(12)
doc.add_paragraph()

# ---- 1. Project Status ----
doc.add_heading('1. Project Status', level=1)
doc.add_paragraph(
    'DeltaRent (game account rental management system) backend is built on '
    'Java 21 + Spring Boot 3.3.5 + MyBatis, with MySQL 8.0 and Spring Security 6 + JWT.'
)
doc.add_paragraph('Completed modules:')
for item in [
    'Domain layer: SysUser, RentalProduct, RentalOrder, Notice, AppealRecord (5 entities)',
    'Mapper layer: 5 mapper interfaces covering all CRUD operations',
    'Service layer: UserService, ProductService, OrderService, NoticeService, AppealService (5 services)',
    'Controller layer: 8 controllers (Auth, Portal, Rental, Order, Notice, Appeal, Dashboard, Health)',
    'Security: JWT auth + BCrypt password encoding + role-based access control (ADMIN/USER)',
    'Database: 6 business tables with seed data',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Issues identified:')
for item in [
    'MySQL service requires Windows admin privileges to start',
    'No automated tests except a single contextLoads() placeholder',
    'JwtAuthFilter public path whitelist inconsistent with SecurityConfig, causing 401 on GET /api/rentals and /api/notices',
    'Role system (sys_role / sys_user_role tables) not wired into application code',
    'Seed data passwords stored in plaintext',
]:
    doc.add_paragraph(item, style='List Bullet')

# ---- 2. Work Completed ----
doc.add_heading('2. Work Completed Today', level=1)

doc.add_heading('2.1 Resolved Development Environment Bottleneck', level=2)
doc.add_paragraph(
    'Added H2 in-memory database (MySQL compatibility mode) to eliminate the need for '
    'MySQL admin privileges during development. Created a dedicated Spring profile '
    '(application-h2.yml) that allows zero-config startup.'
)
doc.add_paragraph('New files:')
for item in [
    'src/main/resources/application-h2.yml - H2 development profile',
    'src/main/resources/schema-h2.sql - H2-compatible schema DDL',
    'src/main/resources/data-h2.sql - H2 seed data',
]:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph('Modified files:')
doc.add_paragraph('build.gradle - Added H2 database dependency (runtimeOnly)', style='List Bullet')

doc.add_heading('2.2 Fixed JwtAuthFilter Public Path Bug', level=2)
doc.add_paragraph(
    'The JwtAuthFilter PUBLIC_PATHS only included /api/health, /api/auth/, and /api/portal/, '
    'but GET /api/rentals and GET /api/notices were also configured as permitAll in SecurityConfig. '
    'Since the filter runs before authorization rules, these GET endpoints returned 401 instead of 200.'
)
doc.add_paragraph(
    'Fix: Replaced simple path prefix matching with HTTP method + path combination logic. '
    'Health, auth, and portal endpoints are public for all methods. '
    'Rental and notice endpoints are public for GET only, requiring authentication for write operations.'
)
doc.add_paragraph('Modified: config/JwtAuthFilter.java')

doc.add_heading('2.3 Established Three-Layer Testing System', level=2)

# Table
table = doc.add_table(rows=4, cols=5)
table.style = 'Light Grid Accent 1'
headers = ['Layer', 'Command', 'What It Tests', 'Time', 'Dependencies']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True

for row_idx, row_data in enumerate([
    ['Unit Tests', './gradlew test', 'Service logic, state machine, validation', '~20s', 'None (H2 in-memory)'],
    ['API Smoke', 'bash scripts/test-api.sh', 'HTTP codes, JSON structure, auth', '~10s', 'Server running'],
    ['Compilation', './gradlew compileJava', 'Syntax, types, dependencies', '~5s', 'None'],
]):
    for col_idx, text in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = text

doc.add_paragraph()

doc.add_heading('Layer 1: Service Unit Tests (21 test cases)', level=3)
doc.add_paragraph('OrderServiceTest (12 cases):')
for item in [
    'Order creation - valid creation, reject rentHours=0, reject rentHours=null, reject missing product, reject unavailable product',
    'State transitions - WAITING_CONFIRM->IN_PROGRESS allowed, WAITING_CONFIRM->CANCELLED allowed, IN_PROGRESS->COMPLETED allowed, COMPLETED->AFTER_SALE allowed, skip-step rejected, invalid status rejected, COMPLETED->CANCELLED rejected, order not found rejected',
]:
    doc.add_paragraph(item, style='List Bullet 2')

doc.add_paragraph('UserServiceTest (9 cases):')
for item in [
    'Login - correct plaintext password, correct BCrypt password, wrong password rejected, user not found rejected, disabled user rejected, blank username rejected',
    'Registration - valid registration, duplicate username rejected, password < 6 chars rejected',
]:
    doc.add_paragraph(item, style='List Bullet 2')

doc.add_heading('Layer 2: Controller MockMvc Tests (6 test cases)', level=3)
for item in [
    'AuthControllerTest (3 cases): login returns token, login fails with error, register succeeds',
    'RentalControllerTest (3 cases): product list, product detail (found), product detail (not found)',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Layer 3: API End-to-End Smoke Test (26 endpoints)', level=3)
doc.add_paragraph(
    'scripts/test-api.sh - Pure bash + curl implementation covering all 26 endpoints across 8 modules. '
    'Automatically logs in, extracts JWT token, and tests protected endpoints with proper authorization. '
    'Result: 26/26 endpoints passed.'
)

# ---- 3. Test Results ----
doc.add_heading('3. Test Execution Results', level=1)
doc.add_paragraph('Command: ./gradlew test')
doc.add_paragraph('Result: BUILD SUCCESSFUL - 29 tests completed, 0 failures')
doc.add_paragraph('Command: bash scripts/test-api.sh')
doc.add_paragraph('Result: 26 passed / 0 failed')

# ---- 4. File Changes ----
doc.add_heading('4. File Change Summary', level=1)

table2 = doc.add_table(rows=11, cols=3)
table2.style = 'Light Grid Accent 1'
for i, h in enumerate(['File', 'Action', 'Description']):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True

files = [
    ['scripts/test-api.sh', 'Added', 'Full API smoke test script'],
    ['src/main/resources/application-h2.yml', 'Added', 'H2 development profile'],
    ['src/main/resources/schema-h2.sql', 'Added', 'H2-compatible schema'],
    ['src/main/resources/data-h2.sql', 'Added', 'H2 seed data'],
    ['src/server/build.gradle', 'Modified', 'Added H2 database dependency'],
    ['config/JwtAuthFilter.java', 'Modified', 'Fixed public path whitelist'],
    ['src/test/.../service/OrderServiceTest.java', 'Added', 'Order service tests (12 cases)'],
    ['src/test/.../service/UserServiceTest.java', 'Added', 'User service tests (9 cases)'],
    ['src/test/.../controller/AuthControllerTest.java', 'Added', 'Auth controller tests (3 cases)'],
    ['src/test/.../controller/RentalControllerTest.java', 'Added', 'Rental controller tests (3 cases)'],
]
for row_idx, row_data in enumerate(files):
    for col_idx, text in enumerate(row_data):
        table2.rows[row_idx + 1].cells[col_idx].text = text

doc.add_paragraph()

# ---- 5. Next Steps ----
doc.add_heading('5. Next Steps', level=1)

doc.add_paragraph('High priority:')
for item in [
    'Role system implementation - wire sys_role / sys_user_role tables into application code',
    'User management API - create /api/admin/users endpoints',
    'BCrypt seed passwords - replace plaintext passwords in seed data',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Medium priority:')
for item in [
    'Operation log - implement domain/mapper/service for operation_log table',
    'Order timeline - populate start_time / end_time during status transitions',
    'Frontend-facing APIs - user profile, change password',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Low priority:')
for item in [
    'Redis integration - token blacklist for logout',
    'Global exception handling - @ControllerAdvice',
    'Bean Validation - replace manual validation with @Valid',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.save('E:/DeltaRent/docs/2026-05-01-backend-report.docx')
print('Document saved to docs/2026-05-01-backend-report.docx')
